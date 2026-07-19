#!/usr/bin/env python3
"""Generate a formatted Word document from paper content and figure screenshots."""

import argparse
import json
import os
import sys
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


DEFAULT_FONT_SIZE = 10.5


def _set_east_asian_font(run, font_name):
    """Set the East-Asian font on a run element."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def set_font(run, text, size=DEFAULT_FONT_SIZE, bold=False):
    """Set font properties: Song Ti for Chinese, Times New Roman for English."""
    run.font.size = Pt(size)
    run.font.bold = bold
    has_chinese = any("\u4e00" <= c <= "\u9fff" for c in text)

    if has_chinese:
        run.font.name = "宋体"
        _set_east_asian_font(run, "宋体")
    else:
        run.font.name = "Times New Roman"
        _set_east_asian_font(run, "宋体")


def apply_body_paragraph_format(p, size=DEFAULT_FONT_SIZE, first_line_indent=True):
    """Apply justified alignment and first-line indent (2 Chinese characters) to body paragraphs."""
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        # 2 Chinese characters ~= 2 * font size in points
        p.paragraph_format.first_line_indent = Pt(size * 2)
    else:
        p.paragraph_format.first_line_indent = Pt(0)


def add_section(doc, item, figure_dir):
    """Add a document element based on the item type."""
    item_type = item.get("type", "paragraph")

    if item_type == "title":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(item["text"])
        set_font(run, item["text"], item.get("size", 16), True)

    elif item_type == "heading":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(item["text"])
        set_font(run, item["text"], item.get("size", DEFAULT_FONT_SIZE), True)

    elif item_type == "subheading":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(item["text"])
        set_font(run, item["text"], item.get("size", DEFAULT_FONT_SIZE), True)

    elif item_type == "paragraph":
        size = item.get("size", DEFAULT_FONT_SIZE)
        p = doc.add_paragraph()
        apply_body_paragraph_format(p, size, item.get("first_line_indent", True))
        run = p.add_run(item["text"])
        set_font(run, item["text"], size, item.get("bold", False))

    elif item_type == "image":
        img_path = os.path.join(figure_dir, item["file"])
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            run.add_picture(img_path, width=Inches(item.get("width", 5.5)))

    elif item_type == "reference":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(item["text"])
        set_font(run, item["text"], item.get("size", DEFAULT_FONT_SIZE), False)

    elif item_type == "source":
        # "文献来源" heading + DOI/URL paragraph (no indent, left-aligned)
        heading_p = doc.add_paragraph()
        heading_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_p.paragraph_format.first_line_indent = Pt(0)
        heading_run = heading_p.add_run("文献来源")
        set_font(heading_run, "文献来源", item.get("size", DEFAULT_FONT_SIZE), True)

        url_p = doc.add_paragraph()
        url_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        url_p.paragraph_format.first_line_indent = Pt(0)
        url_run = url_p.add_run(item["text"])
        set_font(url_run, item["text"], item.get("size", DEFAULT_FONT_SIZE), False)

    elif item_type == "blank":
        doc.add_paragraph()


def main():
    parser = argparse.ArgumentParser(description="Generate a paper-sharing Word document")
    parser.add_argument("--content", help="JSON content config file (omit to read from stdin)")
    parser.add_argument("--figures", required=True, help="Directory containing figure screenshots")
    parser.add_argument("--output", required=True, help="Output path for the Word document")
    parser.add_argument("--cleanup", action="store_true", help="Delete figures directory after generation")
    args = parser.parse_args()

    if args.content:
        with open(args.content, "r", encoding="utf-8") as f:
            content = json.load(f)
    else:
        content = json.load(sys.stdin)

    doc = Document()

    # Set default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(DEFAULT_FONT_SIZE)
    _set_east_asian_font(style, "宋体")

    for item in content.get("sections", []):
        add_section(doc, item, args.figures)

    doc.save(args.output)
    print(f"Word document saved: {args.output}")

    if args.cleanup and os.path.exists(args.figures):
        shutil.rmtree(args.figures)
        print(f"Cleaned up: {args.figures}")


if __name__ == "__main__":
    main()
